import logging
import hashlib
import io
import re
from typing import Dict, Any, Optional, List
from datetime import datetime
import aiohttp

from .config import government_scraper_settings

class GovernmentDocumentProcessor:
    """
    Processor for government documents with text extraction and content analysis
    """

    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.settings = government_scraper_settings
        self.session = None

    async def __aenter__(self):
        """Async context manager entry"""
        timeout = aiohttp.ClientTimeout(total=self.settings.api_timeout)
        self.session = aiohttp.ClientSession(
            timeout=timeout,
            headers={'User-Agent': self.settings.user_agent}
        )
        return self

    async def __aexit__(self, exc_type, exc_val, exc_tb):
        """Async context manager exit"""
        if self.session:
            await self.session.close()

    async def process_document(
        self,
        document_url: str,
        document_data: Optional[bytes] = None,
        content_type: Optional[str] = None
    ) -> Optional[Dict[str, Any]]:
        """
        Process government document

        Args:
            document_url: URL of the document
            document_data: Document content bytes (optional, will download if not provided)
            content_type: Content type of the document

        Returns:
            Processed document data or None if failed
        """
        try:
            # Download document if not provided
            if document_data is None:
                download_result = await self._download_document(document_url)
                if not download_result:
                    return None
                document_data = download_result['content']
                content_type = download_result['content_type']

            # Check file size
            if len(document_data) > self.settings.max_document_size:
                self.logger.warning(f"Document too large: {len(document_data)} bytes")
                return None

            # Extract text content
            text_content = await self._extract_text(document_data, content_type)
            if not text_content:
                return None

            # Extract metadata
            metadata = await self._extract_metadata(document_data, content_type, document_url)

            # Analyze content
            analysis = await self._analyze_content(text_content)

            # Calculate content hash
            content_hash = hashlib.sha256(text_content.encode('utf-8')).hexdigest()

            return {
                'url': document_url,
                'text_content': text_content,
                'metadata': metadata,
                'analysis': analysis,
                'content_hash': content_hash,
                'word_count': len(text_content.split()),
                'character_count': len(text_content),
                'processing_timestamp': datetime.utcnow().isoformat(),
                'source': 'government_document'
            }

        except Exception as e:
            self.logger.error(f"Document processing failed for {document_url}: {e}")
            return None

    async def _download_document(self, document_url: str) -> Optional[Dict[str, Any]]:
        """
        Download document from URL

        Args:
            document_url: URL of the document to download

        Returns:
            Dictionary with content and content_type or None if failed
        """
        try:
            async with self.session.get(document_url) as response:
                if response.status == 200:
                    content = await response.read()
                    content_type = response.headers.get('content-type', '')

                    return {
                        'content': content,
                        'content_type': content_type
                    }
                else:
                    self.logger.error(f"Document download failed: {response.status} - {document_url}")
                    return None

        except Exception as e:
            self.logger.error(f"Document download failed for {document_url}: {e}")
            return None

    async def _extract_text(
        self,
        document_data: bytes,
        content_type: str
    ) -> Optional[str]:
        """
        Extract text from document

        Args:
            document_data: Document content bytes
            content_type: Content type of the document

        Returns:
            Extracted text or None if failed
        """
        try:
            content_type_lower = content_type.lower()

            if 'pdf' in content_type_lower:
                return await self._extract_pdf_text(document_data)
            elif 'word' in content_type_lower or 'docx' in content_type_lower:
                return await self._extract_docx_text(document_data)
            elif 'text' in content_type_lower or 'plain' in content_type_lower:
                return document_data.decode('utf-8', errors='ignore')
            elif 'html' in content_type_lower:
                return await self._extract_html_text(document_data)
            else:
                self.logger.warning(f"Unsupported content type: {content_type}")
                return None

        except Exception as e:
            self.logger.error(f"Text extraction failed: {e}")
            return None

    async def _extract_pdf_text(self, pdf_data: bytes) -> str:
        """
        Extract text from PDF document

        Args:
            pdf_data: PDF document bytes

        Returns:
            Extracted text
        """
        try:
            # Try to use PyPDF2 first
            try:
                import PyPDF2
                text_content = []
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_data))

                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)

                return '\n\n'.join(text_content)

            except ImportError:
                self.logger.warning("PyPDF2 not available, trying alternative method")
                return self._extract_pdf_text_fallback(pdf_data)

        except Exception as e:
            self.logger.error(f"PDF text extraction failed: {e}")
            return ""

    def _extract_pdf_text_fallback(self, pdf_data: bytes) -> str:
        """
        Fallback PDF text extraction method

        Args:
            pdf_data: PDF document bytes

        Returns:
            Extracted text
        """
        try:
            # Simple text extraction using regex patterns
            text = pdf_data.decode('utf-8', errors='ignore')

            # Remove PDF-specific patterns
            text = re.sub(r'%PDF.*?%%EOF', '', text, flags=re.DOTALL)
            text = re.sub(r'[^\x00-\x7F]+', '', text)  # Remove non-ASCII
            text = re.sub(r'\s+', ' ', text)  # Normalize whitespace

            return text.strip()

        except Exception as e:
            self.logger.error(f"PDF fallback extraction failed: {e}")
            return ""

    async def _extract_docx_text(self, docx_data: bytes) -> str:
        """
        Extract text from DOCX document

        Args:
            docx_data: DOCX document bytes

        Returns:
            Extracted text
        """
        try:
            # Try to use python-docx
            try:
                import docx
                doc = docx.Document(io.BytesIO(docx_data))

                text_content = []
                for paragraph in doc.paragraphs:
                    if paragraph.text:
                        text_content.append(paragraph.text)

                return '\n\n'.join(text_content)

            except ImportError:
                self.logger.warning("python-docx not available, trying alternative method")
                return self._extract_docx_text_fallback(docx_data)

        except Exception as e:
            self.logger.error(f"DOCX text extraction failed: {e}")
            return ""

    def _extract_docx_text_fallback(self, docx_data: bytes) -> str:
        """
        Fallback DOCX text extraction method

        Args:
            docx_data: DOCX document bytes

        Returns:
            Extracted text
        """
        try:
            # Simple text extraction for DOCX
            text = docx_data.decode('utf-8', errors='ignore')

            # Remove DOCX-specific patterns
            text = re.sub(r'PK.*?PK', '', text, flags=re.DOTALL)
            text = re.sub(r'[^\x00-\x7F]+', '', text)  # Remove non-ASCII
            text = re.sub(r'\s+', ' ', text)  # Normalize whitespace

            return text.strip()

        except Exception as e:
            self.logger.error(f"DOCX fallback extraction failed: {e}")
            return ""

    async def _extract_html_text(self, html_data: bytes) -> str:
        """
        Extract text from HTML document

        Args:
            html_data: HTML document bytes

        Returns:
            Extracted text
        """
        try:
            from bs4 import BeautifulSoup

            html_content = html_data.decode('utf-8', errors='ignore')
            soup = BeautifulSoup(html_content, 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Get text
            text = soup.get_text()

            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)

            return text

        except Exception as e:
            self.logger.error(f"HTML text extraction failed: {e}")
            return ""

    async def _extract_metadata(
        self,
        document_data: bytes,
        content_type: str,
        document_url: str
    ) -> Dict[str, Any]:
        """
        Extract document metadata

        Args:
            document_data: Document content bytes
            content_type: Content type of the document
            document_url: URL of the document

        Returns:
            Metadata dictionary
        """
        metadata = {
            'content_type': content_type,
            'file_size': len(document_data),
            'url': document_url,
            'extraction_timestamp': datetime.utcnow().isoformat()
        }

        # Add content-type specific metadata
        content_type_lower = content_type.lower()

        if 'pdf' in content_type_lower:
            metadata.update(await self._extract_pdf_metadata(document_data))
        elif 'word' in content_type_lower or 'docx' in content_type_lower:
            metadata.update(await self._extract_docx_metadata(document_data))

        return metadata

    async def _extract_pdf_metadata(self, pdf_data: bytes) -> Dict[str, Any]:
        """
        Extract PDF metadata

        Args:
            pdf_data: PDF document bytes

        Returns:
            PDF metadata dictionary
        """
        try:
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_data))

            metadata = {
                'page_count': len(pdf_reader.pages),
                'pdf_version': pdf_reader.metadata.get('/PDFVersion', 'Unknown'),
                'creator': pdf_reader.metadata.get('/Creator', 'Unknown'),
                'producer': pdf_reader.metadata.get('/Producer', 'Unknown'),
                'creation_date': pdf_reader.metadata.get('/CreationDate', 'Unknown'),
                'modification_date': pdf_reader.metadata.get('/ModDate', 'Unknown')
            }

            return metadata

        except Exception as e:
            self.logger.error(f"PDF metadata extraction failed: {e}")
            return {}

    async def _extract_docx_metadata(self, docx_data: bytes) -> Dict[str, Any]:
        """
        Extract DOCX metadata

        Args:
            docx_data: DOCX document bytes

        Returns:
            DOCX metadata dictionary
        """
        try:
            import docx
            doc = docx.Document(io.BytesIO(docx_data))

            metadata = {
                'paragraph_count': len(doc.paragraphs),
                'section_count': len(doc.sections)
            }

            # Try to extract core properties
            try:
                core_props = doc.core_properties
                if core_props:
                    metadata.update({
                        'title': core_props.title,
                        'author': core_props.author,
                        'subject': core_props.subject,
                        'created': core_props.created,
                        'modified': core_props.modified
                    })
            except Exception:
                pass

            return metadata

        except Exception as e:
            self.logger.error(f"DOCX metadata extraction failed: {e}")
            return {}

    async def _analyze_content(self, text_content: str) -> Dict[str, Any]:
        """
        Analyze document content

        Args:
            text_content: Document text content

        Returns:
            Content analysis dictionary
        """
        try:
            analysis = {
                'language': self._detect_language(text_content),
                'document_type': self._classify_document_type(text_content),
                'key_topics': self._extract_key_topics(text_content),
                'entities': self._extract_entities(text_content),
                'summary': await self._generate_summary(text_content),
                'government_terms': self._extract_government_terms(text_content)
            }

            return analysis

        except Exception as e:
            self.logger.error(f"Content analysis failed: {e}")
            return {}

    def _detect_language(self, text: str) -> str:
        """
        Detect document language

        Args:
            text: Document text

        Returns:
            Language code ('id' for Indonesian, 'en' for English)
        """
        try:
            # Simple language detection based on common words
            indonesian_words = [
                'yang', 'dan', 'atau', 'dalam', 'untuk', 'dengan', 'oleh', 'dari', 'ke', 'di',
                'peraturan', 'pemerintah', 'kementerian', 'republik', 'indonesia', 'negara',
                'undang', 'undang', 'hukum', 'perundangan', 'keputusan', 'surat', 'edaran'
            ]

            english_words = [
                'the', 'and', 'or', 'in', 'for', 'with', 'by', 'from', 'to', 'at',
                'regulation', 'government', 'ministry', 'republic', 'indonesia', 'state',
                'law', 'legal', 'decision', 'letter', 'circular'
            ]

            text_lower = text.lower()
            indonesian_count = sum(1 for word in indonesian_words if word in text_lower)
            english_count = sum(1 for word in english_words if word in text_lower)

            return 'id' if indonesian_count > english_count else 'en'

        except Exception as e:
            self.logger.error(f"Language detection failed: {e}")
            return 'en'

    def _classify_document_type(self, text: str) -> str:
        """
        Classify document type

        Args:
            text: Document text

        Returns:
            Document type classification
        """
        try:
            text_lower = text.lower()

            # Indonesian government document types
            if any(word in text_lower for word in ['peraturan', 'regulation', 'law', 'undang-undang']):
                return 'regulation'
            elif any(word in text_lower for word in ['laporan', 'report', 'annual', 'tahunan']):
                return 'report'
            elif any(word in text_lower for word in ['keputusan', 'decision', 'decree']):
                return 'decision'
            elif any(word in text_lower for word in ['surat', 'letter', 'circular', 'edaran']):
                return 'letter'
            elif any(word in text_lower for word in ['perjanjian', 'agreement', 'treaty']):
                return 'agreement'
            elif any(word in text_lower for word in ['kebijakan', 'policy', 'guideline']):
                return 'policy'
            else:
                return 'document'

        except Exception as e:
            self.logger.error(f"Document type classification failed: {e}")
            return 'document'

    def _extract_key_topics(self, text: str) -> List[str]:
        """
        Extract key topics from document

        Args:
            text: Document text

        Returns:
            List of key topics
        """
        try:
            # Simple keyword extraction based on frequency
            words = re.findall(r'\b\w+\b', text.lower())

            # Remove common stop words
            stop_words = {
                'yang', 'dan', 'atau', 'dalam', 'untuk', 'dengan', 'oleh', 'dari', 'ke', 'di',
                'the', 'and', 'or', 'in', 'for', 'with', 'by', 'from', 'to', 'at', 'is', 'are',
                'a', 'an', 'of', 'on', 'as', 'be', 'have', 'has', 'had', 'do', 'does', 'did'
            }

            # Count word frequency
            word_freq = {}
            for word in words:
                if len(word) > 3 and word not in stop_words:
                    word_freq[word] = word_freq.get(word, 0) + 1

            # Get top keywords
            sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
            return [word for word, freq in sorted_words[:10]]

        except Exception as e:
            self.logger.error(f"Key topics extraction failed: {e}")
            return []

    def _extract_entities(self, text: str) -> List[str]:
        """
        Extract entities from document

        Args:
            text: Document text

        Returns:
            List of entities
        """
        try:
            # Simple entity extraction using regex patterns
            entities = []

            # Government agencies
            agency_pattern = r'\b(?:Kementerian|Ministry|Badan|Agency|Direktorat|Directorate)\s+[A-Z][a-zA-Z\s]+\b'
            agencies = re.findall(agency_pattern, text, re.IGNORECASE)
            entities.extend(agencies)

            # Government officials
            official_pattern = r'\b(?:Menteri|Minister|Direktur|Director|Kepala|Head)\s+[A-Z][a-zA-Z\s]+\b'
            officials = re.findall(official_pattern, text, re.IGNORECASE)
            entities.extend(officials)

            # Dates
            date_pattern = r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b'
            dates = re.findall(date_pattern, text)
            entities.extend(dates)

            return list(set(entities))  # Remove duplicates

        except Exception as e:
            self.logger.error(f"Entity extraction failed: {e}")
            return []

    async def _generate_summary(self, text: str) -> str:
        """
        Generate document summary

        Args:
            text: Document text

        Returns:
            Document summary
        """
        try:
            # Simple summary generation (first few sentences)
            sentences = re.split(r'[.!?]+', text)
            sentences = [s.strip() for s in sentences if s.strip()]

            # Take first 3 sentences as summary
            summary_sentences = sentences[:3]
            summary = '. '.join(summary_sentences)

            # Limit summary length
            if len(summary) > self.settings.max_summary_length:
                summary = summary[:self.settings.max_summary_length] + '...'

            return summary

        except Exception as e:
            self.logger.error(f"Summary generation failed: {e}")
            return ""

    def _extract_government_terms(self, text: str) -> List[str]:
        """
        Extract government-specific terms

        Args:
            text: Document text

        Returns:
            List of government terms
        """
        try:
            government_terms = [
                'pemerintah', 'government', 'kementerian', 'ministry', 'negara', 'state',
                'republik', 'republic', 'indonesia', 'undang-undang', 'law', 'peraturan',
                'regulation', 'keputusan', 'decision', 'surat', 'letter', 'edaran',
                'circular', 'kebijakan', 'policy', 'pedoman', 'guideline', 'standar',
                'standard', 'prosedur', 'procedure', 'perjanjian', 'agreement'
            ]

            found_terms = []
            text_lower = text.lower()

            for term in government_terms:
                if term in text_lower:
                    found_terms.append(term)

            return list(set(found_terms))

        except Exception as e:
            self.logger.error(f"Government terms extraction failed: {e}")
            return []
